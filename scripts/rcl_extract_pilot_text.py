#!/usr/bin/env python3
"""Extract local text and machine observations for the RCL gold pilot."""

from __future__ import annotations

import argparse
import concurrent.futures
import csv
import hashlib
import json
import re
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree


PROTOCOL_ID = "rcl-gold-pilot-text-extraction"
PROTOCOL_VERSION = "0.2.0"

OBSERVATION_FIELDS = [
    "queue_id",
    "project_id",
    "filename",
    "source_media_type",
    "raw_sha256",
    "text_sha256",
    "text_path",
    "extract_status",
    "extractor",
    "page_or_section_count",
    "char_count",
    "word_count",
    "line_count",
    "machine_doc_type_hint",
    "machine_source_type_hint",
    "machine_pii_hint",
    "machine_pii_types",
    "machine_extraction_quality_hint",
    "machine_review_priority",
    "machine_notes",
]

POLISH_WORD_RE = re.compile(r"[0-9A-Za-ząćęłńóśźżĄĆĘŁŃÓŚŹŻ]+")
EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_RE = re.compile(r"(?<!\d)(?:\+?48[\s.-]?)?(?:\d[\s.-]?){9}(?!\d)")
PESEL_RE = re.compile(r"(?<!\d)\d{11}(?!\d)")
ADDRESS_RE = re.compile(
    r"\b(?:ul\.|ulica|al\.|aleja|pl\.|plac)\s+[A-ZĄĆĘŁŃÓŚŹŻ][^\n,;]{2,80}\s+\d",
    re.IGNORECASE,
)
SIGNATURE_RE = re.compile(
    r"\b(?:podpis|podpisano|z poważaniem|z powazaniem|prezes|członek zarządu|"
    r"czlonek zarzadu|dyrektor)\b",
    re.IGNORECASE,
)


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def sha256_bytes(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_jsonl(path: Path) -> list[dict[str, object]]:
    with path.open(encoding="utf-8") as handle:
        return [json.loads(line) for line in handle if line.strip()]


def write_jsonl(path: Path, rows: list[dict[str, object]]) -> None:
    temporary_path = path.with_suffix(path.suffix + ".tmp")
    with temporary_path.open("w", encoding="utf-8", newline="\n") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n")
    temporary_path.replace(path)


def write_observations(path: Path, rows: list[dict[str, str]]) -> None:
    temporary_path = path.with_suffix(path.suffix + ".tmp")
    with temporary_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=OBSERVATION_FIELDS)
        writer.writeheader()
        writer.writerows(rows)
    temporary_path.replace(path)


def normalize_text(parts: list[str]) -> str:
    text = "\n\n".join(part.strip() for part in parts if part and part.strip())
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip() + ("\n" if text.strip() else "")


def extract_pdf(path: Path) -> tuple[str, str, int, list[str]]:
    warnings: list[str] = []
    try:
        import pdfplumber

        pages: list[str] = []
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
        text = normalize_text(pages)
        if text.strip():
            return text, "pdfplumber", page_count, warnings
        warnings.append("pdfplumber returned empty text")
    except Exception as exc:  # pragma: no cover - dependency behavior varies.
        warnings.append(f"pdfplumber failed: {exc}")

    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return normalize_text(pages), "pypdf", len(reader.pages), warnings
    except Exception as exc:  # pragma: no cover - dependency behavior varies.
        warnings.append(f"pypdf failed: {exc}")
        return "", "pdf_failed", 0, warnings


def extract_docx(path: Path) -> tuple[str, str, int, list[str]]:
    warnings: list[str] = []
    try:
        from docx import Document

        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return normalize_text(parts), "python-docx", len(document.paragraphs), warnings
    except Exception as exc:
        warnings.append(f"python-docx failed: {exc}")

    try:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml")
        root = ElementTree.fromstring(xml)
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        parts = [
            "".join(node.itertext())
            for node in root.findall(".//w:p", namespace)
            if "".join(node.itertext()).strip()
        ]
        return normalize_text(parts), "docx-xml", len(parts), warnings
    except Exception as exc:
        warnings.append(f"docx xml fallback failed: {exc}")
        return "", "docx_failed", 0, warnings


def classify_doc_type(filename: str, category: str, text: str) -> str:
    filename_and_category = f"{filename}\n{category}".lower()
    combined = f"{filename_and_category}\n{text[:4000]}".lower()
    if "stanowisk" in filename_and_category:
        return "organization_comment"
    if "odniesienie" in filename_and_category and "uwag" in filename_and_category:
        return "government_response"
    if "pismo kierujące" in filename_and_category or "pismo kierujace" in filename_and_category:
        return "cover_letter"
    if "projekt" in filename_and_category and (
        "ustawy" in filename_and_category or "rozporządzenia" in filename_and_category
    ):
        return "draft_law"
    if "stanowisk" in combined or "uwag" in combined or "opinia" in combined:
        return "organization_comment"
    return "other"


def classify_source_type(filename: str, text: str) -> str:
    head = f"{filename}\n{text[:4000]}".lower()
    if any(term in head for term in ["związek zawodowy", "zwiazek zawodowy", "solidarność"]):
        return "trade_union"
    if any(term in head for term in ["pracodawców", "pracodawcow", "konfederacja lewiatan"]):
        return "employer_organization"
    if any(term in head for term in ["fundacja", "stowarzyszenie", "instytut"]):
        return "ngo"
    if any(term in head for term in ["izba", "samorząd", "samorzad", "rada adwokacka", "lekarska"]):
        return "professional_body"
    if any(term in head for term in ["sp. z o.o.", "spółka", "spolka", " s.a.", " sa "]):
        return "company"
    if any(
        term in head
        for term in ["główny inspektor", "glowny inspektor", "prezes urzędu", "prezes urzedu"]
    ):
        return "public_body"
    if any(term in head for term in ["osoba fizyczna", "obywatel", "obywatelka"]):
        return "individual"
    return "unknown"


def pii_observations(text: str) -> tuple[str, str]:
    hits: list[str] = []
    if EMAIL_RE.search(text):
        hits.append("email")
    if PHONE_RE.search(text):
        hits.append("phone")
    if PESEL_RE.search(text):
        hits.append("personal_identifier")
    if ADDRESS_RE.search(text):
        hits.append("address")
    if SIGNATURE_RE.search(text):
        hits.append("signature_or_named_role")
    if hits:
        return "yes", ";".join(hits)
    return "uncertain", ""


def quality_hint(char_count: int, word_count: int, pages: int, status: str) -> str:
    if status != "extracted" or char_count == 0:
        return "not_extractable"
    if char_count < 500 or word_count < 80:
        return "poor"
    if pages and char_count / pages < 350:
        return "poor"
    if char_count < 1500:
        return "usable"
    return "good"


def review_priority(machine_pii_hint: str, quality: str, doc_type: str) -> str:
    if machine_pii_hint == "yes":
        return "pii_first"
    if quality in {"poor", "not_extractable"}:
        return "extraction_first"
    if doc_type != "organization_comment":
        return "scope_check"
    return "standard"


def extract_one(path: Path, media_type: str) -> tuple[str, str, int, list[str]]:
    suffix = path.suffix.lower()
    if media_type == "application/pdf" or suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx" or media_type.endswith("wordprocessingml.document"):
        return extract_docx(path)
    return "", "unsupported", 0, [f"unsupported media type: {media_type}"]


def extract_manifest_item(
    item: dict[str, object],
    root: Path,
    text_dir: Path,
    overwrite: bool,
    prior_evidence: dict[str, dict[str, object]],
) -> dict[str, object]:
    version = item["version"]
    object_id = str(item["object"]["id"])
    queue_id = object_id.rsplit(":", 1)[-1]
    raw_path = root / str(version["local_path"])
    media_type = str(version["media_type"])
    raw_sha256 = str(version["digest"]["value"])
    text_path = text_dir / f"{queue_id}.txt"

    warnings: list[str] = []
    status = "extracted"
    extractor = ""
    page_count = 0
    if text_path.is_file() and not overwrite:
        text = text_path.read_text(encoding="utf-8", errors="replace")
        evidence = prior_evidence.get(queue_id, {})
        status = str(evidence.get("extract_status", "extracted" if text.strip() else "empty_text"))
        extractor = str(evidence.get("extractor", "existing"))
        page_count = int(evidence.get("page_or_section_count", 0) or 0)
        warnings = [str(value) for value in evidence.get("warnings", [])]
    else:
        if not raw_path.is_file():
            text = ""
            status = "missing_raw"
            extractor = "missing_raw"
            warnings.append(f"missing raw file: {raw_path}")
        elif sha256_file(raw_path) != raw_sha256:
            text = ""
            status = "raw_digest_mismatch"
            extractor = "not_run"
            warnings.append("raw file digest does not match source manifest")
        else:
            text, extractor, page_count, warnings = extract_one(raw_path, media_type)
            if not text.strip():
                status = "empty_text"
            temporary_path = text_path.with_suffix(".txt.tmp")
            temporary_path.write_text(text, encoding="utf-8", newline="\n")
            temporary_path.replace(text_path)

    return {
        "item": item,
        "object_id": object_id,
        "queue_id": queue_id,
        "raw_path": raw_path,
        "media_type": media_type,
        "raw_sha256": raw_sha256,
        "text_path": text_path,
        "text": text,
        "warnings": warnings,
        "status": status,
        "extractor": extractor,
        "page_count": page_count,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract text for the RCL gold-pilot local review set.")
    parser.add_argument("--input-dir", default="data/rcl_gold_pilot_v0_1")
    parser.add_argument("--actor", default="unassigned")
    parser.add_argument("--overwrite", action="store_true")
    parser.add_argument("--checkpoint-every", type=int, default=25)
    parser.add_argument("--workers", type=int, default=1)
    args = parser.parse_args()
    if args.checkpoint_every < 1:
        parser.error("--checkpoint-every must be at least 1")
    if args.workers < 1:
        parser.error("--workers must be at least 1")

    started_at = utc_now()
    run_id = "rcl-extract-pilot-" + started_at.replace(":", "").replace("-", "").replace("+", "z")
    root = Path(args.input_dir)
    manifest_path = root / "source_manifest.jsonl"
    text_dir = root / "extracted_text"
    runs_dir = root / "runs"
    text_dir.mkdir(parents=True, exist_ok=True)
    runs_dir.mkdir(parents=True, exist_ok=True)

    source_manifest = read_jsonl(manifest_path)
    extraction_manifest_path = root / "extraction_manifest.jsonl"
    observations_path = root / "machine_observations.csv"
    prior_evidence: dict[str, dict[str, object]] = {}
    if extraction_manifest_path.exists() and not args.overwrite:
        for entry in read_jsonl(extraction_manifest_path):
            queue_id = str(entry["object"]["id"]).rsplit(":", 1)[-1]
            prior_evidence[queue_id] = dict(entry.get("run_evidence", {}))
    extraction_manifest: list[dict[str, object]] = []
    observations: list[dict[str, str]] = []
    counters: Counter[str] = Counter()

    executor = concurrent.futures.ThreadPoolExecutor(max_workers=args.workers)
    payloads = executor.map(
        lambda item: extract_manifest_item(
            item,
            root,
            text_dir,
            args.overwrite,
            prior_evidence,
        ),
        source_manifest,
    )
    for index, payload in enumerate(payloads, start=1):
        item = payload["item"]
        version = item["version"]
        source = item["source"]
        object_id = str(payload["object_id"])
        queue_id = str(payload["queue_id"])
        raw_path = payload["raw_path"]
        media_type = str(payload["media_type"])
        raw_sha256 = str(payload["raw_sha256"])
        text_path = payload["text_path"]
        text = str(payload["text"])
        warnings = payload["warnings"]
        status = str(payload["status"])
        extractor = str(payload["extractor"])
        page_count = int(payload["page_count"])

        print(f"{index}/{len(source_manifest)} {queue_id}: {raw_path.name}")
        text_raw = text.encode("utf-8")
        text_sha256 = sha256_bytes(text_raw)
        char_count = len(text)
        words = POLISH_WORD_RE.findall(text)
        word_count = len(words)
        line_count = len([line for line in text.splitlines() if line.strip()])
        filename = str(item["object"]["name"])
        category = str(source["category"])
        doc_type = classify_doc_type(filename, category, text)
        source_type = classify_source_type(filename, text)
        pii_hint, pii_types = pii_observations(text)
        quality = quality_hint(char_count, word_count, page_count, status)
        priority = review_priority(pii_hint, quality, doc_type)
        counters[status] += 1
        counters[f"quality:{quality}"] += 1
        counters[f"pii:{pii_hint}"] += 1
        counters[f"doc_type:{doc_type}"] += 1

        relative_text_path = text_path.relative_to(root).as_posix()
        extraction_manifest.append(
            {
                "object": {
                    "id": f"rcl:extracted-text:{queue_id}",
                    "kind": "local_extracted_text",
                    "name": f"{queue_id}.txt",
                },
                "version": {
                    "id": f"sha256:{text_sha256}",
                    "digest": {"algorithm": "sha256", "value": text_sha256},
                    "bytes": len(text_raw),
                    "chars": char_count,
                    "words": word_count,
                    "lines": line_count,
                    "media_type": "text/plain; charset=utf-8",
                    "local_path": relative_text_path,
                },
                "source": {
                    "source_object_id": object_id,
                    "source_version_id": str(version["id"]),
                    "source_local_path": str(version["local_path"]),
                    "source_url": str(source["url"]),
                },
                "relations": [
                    {
                        "predicate": "DERIVED_FROM",
                        "target": {"object_id": object_id, "version_id": str(version["id"])},
                    },
                    {
                        "predicate": "GENERATED_BY",
                        "target": {"protocol_id": PROTOCOL_ID, "run_id": run_id},
                    },
                ],
                "run_evidence": {
                    "extract_status": status,
                    "extractor": extractor,
                    "page_or_section_count": page_count,
                    "warnings": warnings,
                },
                "review_status": {
                    "manual_review": "not_reviewed",
                    "publication_boundary": "local_text_not_published_before_legal_pii_review",
                },
            }
        )
        observations.append(
            {
                "queue_id": queue_id,
                "project_id": str(source["project_id"]),
                "filename": filename,
                "source_media_type": media_type,
                "raw_sha256": raw_sha256,
                "text_sha256": text_sha256,
                "text_path": relative_text_path,
                "extract_status": status,
                "extractor": extractor,
                "page_or_section_count": str(page_count),
                "char_count": str(char_count),
                "word_count": str(word_count),
                "line_count": str(line_count),
                "machine_doc_type_hint": doc_type,
                "machine_source_type_hint": source_type,
                "machine_pii_hint": pii_hint,
                "machine_pii_types": pii_types,
                "machine_extraction_quality_hint": quality,
                "machine_review_priority": priority,
                "machine_notes": "; ".join(warnings),
            }
        )

        if index % args.checkpoint_every == 0:
            write_jsonl(extraction_manifest_path, extraction_manifest)
            write_observations(observations_path, observations)

    executor.shutdown(wait=True)
    write_jsonl(extraction_manifest_path, extraction_manifest)
    write_observations(observations_path, observations)

    finished_at = utc_now()
    run = {
        "object": {"id": "rcl:gold-set-pilot-text-extraction", "kind": "dataset_processing_stage"},
        "version": {
            "id": f"sha256:{sha256_file(extraction_manifest_path)}",
            "manifest": "extraction_manifest.jsonl",
        },
        "protocol": {
            "id": PROTOCOL_ID,
            "version": PROTOCOL_VERSION,
            "tool": "scripts/rcl_extract_pilot_text.py",
            "input_manifest": "source_manifest.jsonl",
        },
        "run": {
            "id": run_id,
            "actor": args.actor,
            "started_at": started_at,
            "finished_at": finished_at,
        },
        "evidence": {
            "source_rows": len(source_manifest),
            "texts_written": len(observations),
            "status_counts": dict(sorted((k, v) for k, v in counters.items() if ":" not in k)),
            "quality_counts": {
                key.split(":", 1)[1]: value
                for key, value in sorted(counters.items())
                if key.startswith("quality:")
            },
            "pii_hint_counts": {
                key.split(":", 1)[1]: value
                for key, value in sorted(counters.items())
                if key.startswith("pii:")
            },
            "doc_type_hint_counts": {
                key.split(":", 1)[1]: value
                for key, value in sorted(counters.items())
                if key.startswith("doc_type:")
            },
        },
        "claims": [],
        "outputs": {
            "extraction_manifest": "extraction_manifest.jsonl",
            "machine_observations": "machine_observations.csv",
            "local_text_dir": "extracted_text/",
        },
        "publication_boundary": {
            "publishable_before_review": [
                "extraction_manifest.jsonl without extracted text payload",
                "machine_observations.csv without snippets",
                "run summary",
            ],
            "local_only_before_review": ["extracted_text/"],
        },
    }
    run_path = runs_dir / f"{run_id}.json"
    run_path.write_text(json.dumps(run, ensure_ascii=False, indent=2, sort_keys=True), encoding="utf-8")

    print(f"texts_written: {len(observations)}")
    print(f"status_counts: {dict(sorted((k, v) for k, v in counters.items() if ':' not in k))}")
    print(f"quality_counts: {run['evidence']['quality_counts']}")
    print(f"pii_hint_counts: {run['evidence']['pii_hint_counts']}")
    print(f"extraction_manifest_sha256: {sha256_file(extraction_manifest_path)}")
    print(f"output: {root.resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
